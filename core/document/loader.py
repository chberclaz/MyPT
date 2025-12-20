"""
Document loader for various file formats.

Supports:
- Plain text (.txt, .text)
- Markdown (.md, .markdown)
- PDF (.pdf) - requires PyMuPDF: pip install PyMuPDF
- Word (.docx) - requires python-docx: pip install python-docx

Extensible via DocumentLoader.register_handler() for custom formats.
"""

import os
import logging
from pathlib import Path
from dataclasses import dataclass, field
from typing import List, Dict, Callable, Optional, Iterator
from datetime import datetime

logger = logging.getLogger("mypt.document")

# =============================================================================
# Optional dependency checks
# =============================================================================

# PyMuPDF for PDF support
PYMUPDF_AVAILABLE = False
try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    fitz = None

# python-docx for DOCX support
DOCX_AVAILABLE = False
try:
    from docx import Document as DocxDocument
    from docx.opc.exceptions import PackageNotFoundError
    DOCX_AVAILABLE = True
except ImportError:
    DocxDocument = None
    PackageNotFoundError = Exception


def get_supported_formats() -> Dict[str, bool]:
    """Return dict of formats and their availability."""
    return {
        '.txt': True,
        '.text': True,
        '.md': True,
        '.markdown': True,
        '.pdf': PYMUPDF_AVAILABLE,
        '.docx': DOCX_AVAILABLE,
    }


@dataclass
class Document:
    """
    Represents a loaded document with metadata.
    
    Attributes:
        text: Full document text content
        source: Source file path
        filename: Just the filename (for display)
        format: File format (txt, md, pdf, etc.)
        metadata: Additional metadata (modified time, size, etc.)
    """
    text: str
    source: str
    filename: str
    format: str
    metadata: Dict = field(default_factory=dict)
    
    @property
    def num_chars(self) -> int:
        return len(self.text)
    
    @property
    def num_lines(self) -> int:
        return self.text.count('\n') + 1


# Type for format handlers: (filepath) -> text
FormatHandler = Callable[[str], str]


class DocumentLoader:
    """
    Load documents from files with format-specific handling.
    
    Supported formats:
        - .txt, .text: Plain text files
        - .md, .markdown: Markdown files
        - .pdf: PDF files (requires PyMuPDF)
        - .docx: Word documents (requires python-docx)
    
    Usage:
        loader = DocumentLoader()
        docs = loader.load_directory("workspace/docs")
        
        for doc in docs:
            print(f"{doc.filename}: {doc.num_chars} chars")
    
    Check available formats:
        from core.document.loader import get_supported_formats
        print(get_supported_formats())
    """
    
    # Default supported extensions (text-based, always available)
    SUPPORTED_EXTENSIONS = {'.txt', '.md', '.text', '.markdown'}
    
    def __init__(self):
        self._handlers: Dict[str, FormatHandler] = {
            '.txt': self._load_text,
            '.text': self._load_text,
            '.md': self._load_markdown,
            '.markdown': self._load_markdown,
        }
        
        # Register PDF handler if PyMuPDF is available
        if PYMUPDF_AVAILABLE:
            self._handlers['.pdf'] = self._load_pdf
            self.SUPPORTED_EXTENSIONS = self.SUPPORTED_EXTENSIONS | {'.pdf'}
            logger.debug("PDF support enabled (PyMuPDF)")
        
        # Register DOCX handler if python-docx is available
        if DOCX_AVAILABLE:
            self._handlers['.docx'] = self._load_docx
            self.SUPPORTED_EXTENSIONS = self.SUPPORTED_EXTENSIONS | {'.docx'}
            logger.debug("DOCX support enabled (python-docx)")
    
    def register_handler(self, extension: str, handler: FormatHandler) -> None:
        """
        Register a custom handler for a file extension.
        
        Args:
            extension: File extension including dot (e.g., '.pdf')
            handler: Function that takes filepath and returns text
        """
        self._handlers[extension.lower()] = handler
    
    def _load_text(self, filepath: str) -> str:
        """Load plain text file."""
        with open(filepath, 'r', encoding='utf-8', errors='replace') as f:
            return f.read()
    
    def _load_markdown(self, filepath: str) -> str:
        """Load markdown file (currently same as text, could strip formatting)."""
        with open(filepath, 'r', encoding='utf-8', errors='replace') as f:
            return f.read()
    
    def _load_pdf(self, filepath: str) -> str:
        """
        Load PDF file using PyMuPDF.
        
        Extracts text from all pages, preserving paragraph structure.
        Handles multi-column layouts and embedded fonts.
        """
        if not PYMUPDF_AVAILABLE:
            raise ImportError("PyMuPDF not installed. Install with: pip install PyMuPDF")
        
        try:
            doc = fitz.open(filepath)
            text_parts = []
            
            for page_num, page in enumerate(doc, 1):
                # Extract text with layout preservation
                page_text = page.get_text("text")
                if page_text.strip():
                    text_parts.append(page_text)
            
            doc.close()
            
            # Join pages with double newline
            full_text = "\n\n".join(text_parts)
            
            # Clean up excessive whitespace while preserving paragraphs
            lines = []
            for line in full_text.split('\n'):
                stripped = line.strip()
                if stripped:
                    lines.append(stripped)
                elif lines and lines[-1] != '':
                    lines.append('')  # Preserve paragraph breaks
            
            return '\n'.join(lines)
            
        except Exception as e:
            logger.error(f"Failed to load PDF {filepath}: {e}")
            raise
    
    def _load_docx(self, filepath: str) -> str:
        """
        Load DOCX file using python-docx.
        
        Extracts text from paragraphs and tables.
        Preserves document structure with paragraph breaks.
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx not installed. Install with: pip install python-docx")
        
        try:
            doc = DocxDocument(filepath)
            text_parts = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    text_parts.append(text)
            
            # Extract tables
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        table_text.append(" | ".join(row_text))
                if table_text:
                    text_parts.append("\n".join(table_text))
            
            return "\n\n".join(text_parts)
            
        except PackageNotFoundError:
            logger.error(f"Invalid or corrupted DOCX file: {filepath}")
            raise
        except Exception as e:
            logger.error(f"Failed to load DOCX {filepath}: {e}")
            raise
    
    def _get_file_metadata(self, filepath: str) -> Dict:
        """Extract file metadata."""
        path = Path(filepath)
        stat = path.stat()
        return {
            'size_bytes': stat.st_size,
            'modified_time': datetime.fromtimestamp(stat.st_mtime).isoformat(),
            'created_time': datetime.fromtimestamp(stat.st_ctime).isoformat(),
        }
    
    def load_file(self, filepath: str) -> Optional[Document]:
        """
        Load a single document from file.
        
        Args:
            filepath: Path to file
            
        Returns:
            Document instance or None if format not supported
        """
        path = Path(filepath)
        ext = path.suffix.lower()
        
        if ext not in self._handlers:
            return None
        
        try:
            text = self._handlers[ext](filepath)
            metadata = self._get_file_metadata(filepath)
            
            return Document(
                text=text,
                source=str(path.absolute()),
                filename=path.name,
                format=ext.lstrip('.'),
                metadata=metadata,
            )
        except Exception as e:
            print(f"Error loading {filepath}: {e}")
            return None
    
    def load_directory(
        self, 
        directory: str, 
        recursive: bool = True,
        extensions: Optional[set] = None
    ) -> List[Document]:
        """
        Load all supported documents from a directory.
        
        Args:
            directory: Directory path
            recursive: Whether to search subdirectories
            extensions: Specific extensions to load (default: all supported)
            
        Returns:
            List of Document instances
        """
        docs = []
        exts = extensions or set(self._handlers.keys())
        
        path = Path(directory)
        if not path.is_dir():
            raise ValueError(f"Not a directory: {directory}")
        
        pattern = '**/*' if recursive else '*'
        
        for file_path in path.glob(pattern):
            if file_path.is_file() and file_path.suffix.lower() in exts:
                doc = self.load_file(str(file_path))
                if doc:
                    docs.append(doc)
        
        return docs
    
    def iter_directory(
        self, 
        directory: str, 
        recursive: bool = True
    ) -> Iterator[Document]:
        """
        Iterate over documents in a directory (memory-efficient).
        
        Yields:
            Document instances one at a time
        """
        path = Path(directory)
        pattern = '**/*' if recursive else '*'
        
        for file_path in path.glob(pattern):
            if file_path.is_file() and file_path.suffix.lower() in self._handlers:
                doc = self.load_file(str(file_path))
                if doc:
                    yield doc

